import requests
from requests.auth import HTTPBasicAuth
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from io import BytesIO
from datetime import datetime
import os


pat = "egl7ubikelpdkykekasz64jezo7aplrw3hvsojl7h2kb7c2jcgja"
org_url = "https://tfs.alliancewebpos.com/tfs/WebPOSCollection/WebPOS"
template_path = "Report Template.docx"

my_display_name = "Jacob Barrientos"
my_unique_name = "jacob.barrientos@asi-dev1.com"

def clear_console():
    # For Windows
    if os.name == "nt":
        os.system("cls")
    # For Mac/Linux
    else:
        os.system("clear")

def safe_get_json(url):
    response = requests.get(url, auth=HTTPBasicAuth("", pat))
    if response.status_code != 200:
        try:
            error_data = response.json()
        except Exception:
            error_data = response.text
        print(f"Error fetching URL ({url}):", response.status_code)
        print("Server response:\n", error_data)
        return None
    try:
        return response.json()
    except Exception:
        print(f"Error parsing JSON from URL ({url}).")
        print("Server response:\n", response.text)
        return None

def createReport():
    work_item_id = input("Enter work item ID: ").strip()
    client_name = input("Enter client name: ").strip()

    # ===== Select location type =====
    print("Select location type:")
    print("1: Onsite")
    print("2: Offsite/Remote")

    while True:
        choice = input("Enter 1 or 2: ").strip()
        if choice == "1":
            location_type1, location_type2 = "X", ""
            break
        elif choice == "2":
            location_type1, location_type2 = "", "X"
            break
        else:
            print("Invalid input. Please enter 1 or 2.")

    print(f"Selected: OS='{location_type1}', FS='{location_type2}'")

    output_path = f"Ticket No. {work_item_id} - Status Report for {client_name}.docx"

    # ===== Fetch work item details =====
    work_item_url = f"{org_url}/_apis/wit/workitems/{work_item_id}?api-version=5.1"
    work_item_data = safe_get_json(work_item_url)
    if not work_item_data or "fields" not in work_item_data:
        print(f"❌ Work item ID '{work_item_id}' does not exist.")
        return

    ticket_number = work_item_data.get("id")
    ticket_subject = work_item_data.get("fields", {}).get("System.Title", "")

    # ===== Fetch latest comments =====
    comments_url = f"{org_url}/_apis/wit/workItems/{work_item_id}/comments?api-version=5.1-preview.3"
    comments_data = safe_get_json(comments_url)
    if not comments_data or "comments" not in comments_data:
        print(f"❌ No comments found for work item ID '{work_item_id}'.")
        return

    my_comments = [
        c for c in comments_data.get("comments", [])
        if c.get("createdBy", {}).get("displayName") == my_display_name
        or c.get("createdBy", {}).get("uniqueName") == my_unique_name
    ]
    if not my_comments:
        print("❌ No comments found for your account.")
        return

    # ===== Pick latest comment =====
    my_comments.sort(key=lambda x: x.get("createdDate", ""))
    latest_comment = my_comments[-1]

    # ===== Convert createdDate to MM/DD/YYYY =====
    iso_date = latest_comment.get("createdDate", "")
    comment_date = ""
    if iso_date:
        dt = datetime.strptime(iso_date, "%Y-%m-%dT%H:%M:%S.%fZ")
        comment_date = dt.strftime("%m/%d/%Y")

    raw_html = latest_comment.get("text", "")

    # ===== Helper function to extract fields =====
    def extract_field_multiline(text, field_name, other_labels):
        lines = text.splitlines()
        capture, result_lines = False, []
        for line in lines:
            if capture:
                if any(line.strip().lower().startswith(label.lower() + ":") for label in other_labels):
                    break
                result_lines.append(line)
            elif line.strip().lower() == field_name.lower() + ":":
                capture = True
        return "\n".join(result_lines).strip()

    soup_for_text = BeautifulSoup(raw_html or "", "html.parser")
    clean_text = soup_for_text.get_text("\n", strip=True)

    root_cause = extract_field_multiline(clean_text, "Root Cause", ["Preventive Action", "Next Step", "Status"])
    preventive_action = extract_field_multiline(clean_text, "Preventive Action", ["Root Cause", "Next Step", "Status"])
    next_step = extract_field_multiline(clean_text, "Next Step", ["Root Cause", "Preventive Action", "Status"])
    status = extract_field_multiline(clean_text, "Status", ["Root Cause", "Preventive Action", "Next Step"])

    # ===== Remove extracted fields from main content =====
    final_text = clean_text
    for field in ["Root Cause", "Preventive Action", "Next Step", "Status"]:
        lines, final_lines, skip = final_text.splitlines(), [], False
        for line in lines:
            if line.strip().lower() == field.lower() + ":":
                skip = True
                continue
            if skip and any(l.strip().lower().endswith(":") for l in lines):
                skip = False
            if not skip:
                final_lines.append(line)
        final_text = "\n".join(final_lines)

    final_text_lines = []
    for line in final_text.splitlines():
        stripped = line.strip().rstrip(".")
        if stripped and stripped.upper() not in ("N/A", "PENDING"):
            final_text_lines.append(line)
    final_text = "\n".join(final_text_lines)

    # ===== Function to add HTML content (text + images) =====
    def add_html_content_to_paragraph(paragraph, html_content):
        soup = BeautifulSoup(html_content, "html.parser")
        def recurse(node, para):
            if isinstance(node, NavigableString):
                for line in str(node).splitlines():
                    para.add_run(line)
                    para.add_run().add_break()
            elif isinstance(node, Tag):
                if node.name == "br":
                    para.add_run().add_break()
                elif node.name == "p":
                    for child in node.children:
                        recurse(child, para)
                    para.add_run().add_break()
                elif node.name == "li":
                    para.add_run("• " if node.parent.name != "ul" else "")
                    for child in node.children:
                        recurse(child, para)
                    para.add_run().add_break()
                elif node.name == "img":
                    img_url = node.get("src")
                    if img_url:
                        img_resp = requests.get(img_url, auth=HTTPBasicAuth("", pat))
                        if img_resp.status_code == 200:
                            para.add_run().add_picture(BytesIO(img_resp.content))
                else:
                    for child in node.children:
                        recurse(child, para)
        for child in soup.children:
            recurse(child, paragraph)

    # ===== Inject into DOCX =====
    doc = Document(template_path)
    placeholder_map = {
        "{{TICKETNUM}}": str(ticket_number),
        "{{DATE}}": comment_date,
        "{{TICKETCONTENT}}": ticket_subject,
        "{{RCA}}": root_cause,
        "{{PREVAC}}": preventive_action,
        "{{NEXT}}": next_step,
        "{{CURSTAT}}": status,
        "{{REPORT_CONTENT}}": raw_html,
        "{{OS}}": location_type1,
        "{{FS}}": location_type2,
        "{{CLIENT}}": client_name
    }

    found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for ph in ["{{TICKETNUM}}", "{{TICKETCONTENT}}", "{{DATE}}", "{{OS}}", "{{FS}}", "{{CLIENT}}"]:
                        if ph in paragraph.text:
                            paragraph.text = paragraph.text.replace(ph, placeholder_map[ph])
                            found = True
                    for ph in ["{{RCA}}", "{{PREVAC}}", "{{NEXT}}", "{{CURSTAT}}"]:
                        if ph in paragraph.text:
                            paragraph.text = paragraph.text.replace(ph, placeholder_map[ph])
                            found = True
                    if "{{REPORT_CONTENT}}" in paragraph.text:
                        for run in paragraph.runs:
                            run.text = ""
                        add_html_content_to_paragraph(paragraph, placeholder_map["{{REPORT_CONTENT}}"])
                        found = True

    if found:
        doc.save(output_path)
        print(f"✅ Report saved as {output_path}")
    else:
        print("❌ No placeholders found in the template.")

# ===== Run the report creation =====
while True:
    clear_console()      # clears screen before running
    createReport()       # run the report

    again = input("\nDo you want to create another report? (y/n): ").strip().lower()
    if again != "y":
        print("Exiting...")
        break
