from flask import Flask, render_template, request, send_file, flash
import requests
from requests.auth import HTTPBasicAuth
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from io import BytesIO
from datetime import datetime
import os

app = Flask(__name__)
app.secret_key = "secret123"  # needed for flashing messages

# ===== Your existing settings =====
pat = "egl7ubikelpdkykekasz64jezo7aplrw3hvsojl7h2kb7c2jcgja"
org_url = "https://tfs.alliancewebpos.com/tfs/WebPOSCollection/WebPOS"
template_path = "Report Template.docx"

my_display_name = "Jacob Barrientos"
my_unique_name = "jacob.barrientos@asi-dev1.com"

# ===== Helper functions =====
def safe_get_json(url):
    response = requests.get(url, auth=HTTPBasicAuth("", pat))
    if response.status_code != 200:
        return None
    try:
        return response.json()
    except Exception:
        return None

def extract_field_multiline(text, field_name, other_labels):
    lines = text.splitlines()
    capture, result_lines = False, []
    for line in lines:
        if capture:
            if any(line.strip().lower().startswith(label.lower() + ":") for label in other_labels):
                break
            result_lines.append(line)
        elif line.strip().lower() == field_name.lower() + ":":
            capture = True
    return "\n".join(result_lines).strip()

def add_html_content_to_paragraph(paragraph, html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    def recurse(node, para):
        if isinstance(node, NavigableString):
            for line in str(node).splitlines():
                para.add_run(line)
                para.add_run().add_break()
        elif isinstance(node, Tag):
            if node.name == "br":
                para.add_run().add_break()
            elif node.name == "p":
                for child in node.children:
                    recurse(child, para)
                para.add_run().add_break()
            elif node.name == "li":
                para.add_run("• ")
                for child in node.children:
                    recurse(child, para)
                para.add_run().add_break()
            elif node.name == "img":
                img_url = node.get("src")
                if img_url:
                    img_resp = requests.get(img_url, auth=HTTPBasicAuth("", pat))
                    if img_resp.status_code == 200:
                        para.add_run().add_picture(BytesIO(img_resp.content))
            else:
                for child in node.children:
                    recurse(child, para)
    for child in soup.children:
        recurse(child, paragraph)

# ===== Main route =====
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        work_item_id = request.form.get("work_item_id")
        client_name = request.form.get("client_name")
        location_type = request.form.get("location_type")

        # Location type handling
        location_type1, location_type2 = ("X", "") if location_type == "onsite" else ("", "X")

        # Fetch work item
        work_item_url = f"{org_url}/_apis/wit/workitems/{work_item_id}?api-version=5.1"
        work_item_data = safe_get_json(work_item_url)
        if not work_item_data or "fields" not in work_item_data:
            flash(f"❌ Work item ID '{work_item_id}' does not exist.")
            return render_template("index.html")

        ticket_number = work_item_data.get("id")
        ticket_subject = work_item_data.get("fields", {}).get("System.Title", "")

        # Fetch comments
        comments_url = f"{org_url}/_apis/wit/workItems/{work_item_id}/comments?api-version=5.1-preview.3"
        comments_data = safe_get_json(comments_url)
        if not comments_data or "comments" not in comments_data:
            flash("❌ No comments found.")
            return render_template("index.html")

        my_comments = [
            c for c in comments_data.get("comments", [])
            if c.get("createdBy", {}).get("displayName") == my_display_name
            or c.get("createdBy", {}).get("uniqueName") == my_unique_name
        ]
        if not my_comments:
            flash("❌ No comments found for your account.")
            return render_template("index.html")

        my_comments.sort(key=lambda x: x.get("createdDate", ""))
        latest_comment = my_comments[-1]

        iso_date = latest_comment.get("createdDate", "")
        comment_date = ""
        if iso_date:
            dt = datetime.strptime(iso_date, "%Y-%m-%dT%H:%M:%S.%fZ")
            comment_date = dt.strftime("%m/%d/%Y")

        raw_html = latest_comment.get("text", "")
        soup_for_text = BeautifulSoup(raw_html or "", "html.parser")
        clean_text = soup_for_text.get_text("\n", strip=True)

        root_cause = extract_field_multiline(clean_text, "Root Cause", ["Preventive Action", "Next Step", "Status"])
        preventive_action = extract_field_multiline(clean_text, "Preventive Action", ["Root Cause", "Next Step", "Status"])
        next_step = extract_field_multiline(clean_text, "Next Step", ["Root Cause", "Preventive Action", "Status"])
        status = extract_field_multiline(clean_text, "Status", ["Root Cause", "Preventive Action", "Next Step"])

        # ===== Create DOCX =====
        doc = Document(template_path)
        placeholder_map = {
            "{{TICKETNUM}}": str(ticket_number),
            "{{DATE}}": comment_date,
            "{{TICKETCONTENT}}": ticket_subject,
            "{{RCA}}": root_cause,
            "{{PREVAC}}": preventive_action,
            "{{NEXT}}": next_step,
            "{{CURSTAT}}": status,
            "{{REPORT_CONTENT}}": raw_html,
            "{{OS}}": location_type1,
            "{{FS}}": location_type2,
            "{{CLIENT}}": client_name
        }

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for ph in ["{{TICKETNUM}}", "{{TICKETCONTENT}}", "{{DATE}}", "{{OS}}", "{{FS}}", "{{CLIENT}}"]:
                            if ph in paragraph.text:
                                paragraph.text = paragraph.text.replace(ph, placeholder_map[ph])
                        for ph in ["{{RCA}}", "{{PREVAC}}", "{{NEXT}}", "{{CURSTAT}}"]:
                            if ph in paragraph.text:
                                paragraph.text = paragraph.text.replace(ph, placeholder_map[ph])
                        if "{{REPORT_CONTENT}}" in paragraph.text:
                            for run in paragraph.runs:
                                run.text = ""
                            add_html_content_to_paragraph(paragraph, placeholder_map["{{REPORT_CONTENT}}"])

        output_path = f"Ticket No. {ticket_number} - Status Report for {client_name}.docx"
        doc.save(output_path)

        return send_file(output_path, as_attachment=True)

    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
